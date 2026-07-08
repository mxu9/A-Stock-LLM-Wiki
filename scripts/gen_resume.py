"""Generate a .docx resume from 简历.md"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ---- Helper functions ----
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=10.5, alignment=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_rich_para(segments, space_after=4):
    """segments is a list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for seg in segments:
        if len(seg) == 2:
            text, bold = seg
            italic = False
        else:
            text, bold, italic = seg
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
        run.bold = bold
        run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    return p

def add_rich_bullet(segments, level=0):
    """segments: list of (text, bold) tuples"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    for text, bold in segments:
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        run.bold = bold
    return p

def set_cell_font(cell, text, bold=False, size=9.5):
    """Set cell text with font"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
run = title.add_run('徐  敏')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(22)
run.bold = True

# Personal info
add_para('意向岗位：推理框架性能优化工程师（sglang方向）  |  工作年限：20+年系统底层开发 + 近1年AI工程化落地  |  出生年月：1975年11月',
         size=9.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# ============================================================
# 教育背景
# ============================================================
add_heading_styled('教育背景', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table.rows[0].cells[0], '时间', bold=True)
set_cell_font(table.rows[0].cells[1], '学校', bold=True)
set_cell_font(table.rows[0].cells[2], '专业', bold=True)
set_cell_font(table.rows[1].cells[0], '1994.09 – 1997.07')
set_cell_font(table.rows[1].cells[1], '上海大学')
set_cell_font(table.rows[1].cells[2], '通信工程')
set_cell_font(table.rows[2].cells[0], '2003.09 – 2006.07')
set_cell_font(table.rows[2].cells[1], '上海交通大学')
set_cell_font(table.rows[2].cells[2], '芯片设计（硕士）')

doc.add_paragraph()  # spacer

# ============================================================
# 工作经历
# ============================================================
add_heading_styled('工作经历', level=2)

add_rich_para([('英特尔亚太研发有限公司 | 资深软件工程师', True)], space_after=0)
add_para('2006.02 – 2025.07', size=9, space_after=2)
add_bullet('长期负责Intel机密计算（TDX）核心固件与安全服务架构设计，深耕X86体系结构、内存管理、PCIe I/O及高并发服务优化，具备从硬件底层到高并发服务架构的全栈系统能力。')
add_bullet('近一年全面转型AI工程化方向：独立完成本地双卡（多GPU）推理环境搭建与部署，深入实践大模型后训练（Post-Training）及AI Agent开发，精通Python，具备将底层系统优化方法论赋能AI推理框架的实战经验。')

doc.add_paragraph()

add_rich_para([('上海贝尔 | 软件工程师', True)], space_after=0)
add_para('1998.07 – 2002.10', size=9, space_after=2)
add_bullet('参与通信设备嵌入式系统开发，积累扎实的C/C++底层编程与实时系统调试能力。')

# ============================================================
# 核心项目经验
# ============================================================
add_heading_styled('核心项目经验', level=2)

# --- Project 1 ---
add_rich_para([('1. TDVF（Intel TDX Virtual Firmware）—— 虚拟化启动性能极致优化', True)], space_after=0)
add_para('核心贡献：通过架构级重构，将TD机密虚拟机启动时间优化至与普通虚拟机持平', size=9, space_after=4)
add_bullet('项目背景：TDX虚拟机因安全机制要求，所有Private Memory必须经过"Accept"操作后才能被Guest访问。原始方案在启动过程中同步Accept全部内存，导致严重启动延迟。')
add_rich_bullet([('优化策略一 —— Lazy Accept + 多核并行：', True)])
add_bullet('深入分析TD-Guest启动阶段的内存访问模式，识别出大量内存在启动初期并不需要立即访问；', level=1)
add_bullet('重构内存加载流程，引入Lazy Accept机制——将内存Accept操作从"启动时全量同步执行"改为"按需触发、首次访问时再Accept"；', level=1)
add_bullet('结合多Core并行处理，将可并发的Accept任务分散到多个CPU核心上执行，充分利用多核算力；', level=1)
add_bullet('效果：大幅减少启动路径上的同步阻塞点，单此项优化贡献了启动时间缩短的主要收益。', level=1)
add_rich_bullet([('优化策略二 —— DMA-MMIO Pre-allocation：', True)])
add_bullet('原始DMA-MMIO初始化流程采用"Allocate → Accept → Deallocate"模式，频繁的内存分配/释放/安全接受操作在启动阶段造成大量冗余开销；', level=1)
add_bullet('引入Pre-allocation机制，在初始化阶段一次性预分配并Accept好所需的DMA缓冲池，后续DMA操作直接复用池中内存，消除反复分配/释放/Accept的开销；', level=1)
add_bullet('效果：DMA初始化阶段延迟显著降低，整体启动时间进一步压缩。', level=1)
add_rich_bullet([('最终成果：', True), ('TD-Guest启动时间缩短40%+，性能达到与Non-TD Guest基本持平的水平。', False)])
add_bullet('技术栈：C，X86汇编，EDK II框架，多核并行编程，内存管理，PCIe/DMA')
add_para('项目链接：https://github.com/tianocore/edk2/blob/master/OvmfPkg/IntelTdx/README.md', size=9)

doc.add_paragraph()

# --- Project 2 ---
add_rich_para([('2. vTPM（虚拟TPM服务）—— 多实例架构升级与Rust高性能重构', True)], space_after=0)
add_para('核心贡献：将单实例开源TPM库改造为高并发多租户服务，并用Rust实现极致性能优化', size=9, space_after=4)
add_bullet('项目背景：基于ms-tpm-20-ref开源参考实现开发vtpm-td服务，为TDX平台上所有VM提供TPM 2.0功能。原始ms-tpm-20-ref为单实例串行设计，全局变量散落各处，无法同时服务多个TD-Guest的并发TPM请求。')
add_rich_bullet([('架构重构 —— 多实例支持：', True)])
add_bullet('引入TD-Guest-ID机制：为每个TPM指令请求标记所属的TD-Guest ID，实现请求来源的精准识别与隔离；', level=1)
add_bullet('引入Context抽象：将原本散落在全局命名空间中的所有状态变量封装为Per-Guest Context结构体，每个TD-Guest拥有独立的Context实例；', level=1)
add_bullet('请求路由与隔离：设计请求分发层，根据TD-Guest-ID路由至对应的Context执行TPM命令处理，实现单进程服务多租户的架构；', level=1)
add_bullet('成果：一个vtpm-td进程可同时服务可配置数量的TD-Guest TPM请求，资源隔离完备，无串行阻塞。', level=1)
add_rich_bullet([('性能优化 —— Rust重构：', True)])
add_bullet('将核心TPM命令处理路径用Rust重新实现，利用Rust的零成本抽象、所有权机制和无畏并发（Fearless Concurrency）特性，在保证内存安全的前提下实现极致性能；', level=1)
add_bullet('结合Rust的异步运行时（Tokio），将TPM命令处理进一步优化为高效的事件驱动模型，充分利用多核CPU；', level=1)
add_bullet('成果：多VM并发场景下TPM命令处理吞吐量提升5倍以上，P99延迟降低至原有1/3，同时消除了因全局变量导致的内存安全问题。', level=1)
add_bullet('技术栈：Rust（核心服务），C/C++（与开源库集成），Tokio异步运行时，Epoll事件驱动')

doc.add_paragraph()

# --- Project 3 ---
add_rich_para([('3. TEEIO-Validation —— PCIe设备高吞吐压测与互通性验证', True)], space_after=0)
add_para('核心贡献：开发PCIe设备标准合规性测试套件，验证多厂商设备极端I/O压力下的稳定性', size=9, space_after=4)
add_bullet('设计并实现基于SPDM/TDISP协议的PCIe设备测试框架，模拟高带宽加密/解密及DMA大数据块传输场景，对设备进行满带宽压测。')
add_bullet('深度分析PCIe拓扑结构与Root Port配置，定位并解决多设备同时DMA时的总线仲裁冲突问题，保障澜起科技、Samsung等第三方设备顺利通过认证。')
add_bullet('技术栈：C，PCIe规范，DMA编程，Linux内核驱动接口')

doc.add_paragraph()

# --- Project 4 ---
add_rich_para([('4. AI推理与Agent开发（近期实战）', True)], space_after=8)

# 4.1
add_rich_para([('审计自动化 —— 验证码/文档/合同/票据识别与信息提取', True)], space_after=0)
add_para('核心贡献：构建AI驱动的审计流程自动化系统，大幅提升识别效率并降低Token消耗', size=9, space_after=4)
add_bullet('项目背景：服务于审计流程自动化，需从验证码、PDF合同、票据等非结构化文档中精准提取关键信息，替代人工录入。')
add_rich_bullet([('优化策略：', True)])
add_bullet('验证码预处理：针对不同类型验证码（扭曲字符、干扰背景等），设计图像预处理Pipeline（灰度化/二值化/降噪/分割），将预处理后的清晰图像送交大模型识别，识别率提升至95%以上；', level=1)
add_bullet('PDF智能截取：针对长文档（合同/报告），不整篇送交大模型，而是基于关键词定位和版面分析，精准截取相关页面后再送交识别，大幅减少无效Token消耗；', level=1)
add_bullet('Playwright自动化操作：利用Playwright实现网页端验证码获取、表单填写、结果提交等流程的端到端自动化，全过程无需人工干预。', level=1)
add_rich_bullet([('最终成果：', True), ('识别率大幅提升，Token消耗显著降低，人力投入减少70%+。', False)])
add_bullet('技术栈：Python，OpenCV（图像预处理），PyMuPDF/PDFPlumber，Playwright，大模型API（GPT-4V/Claude Vision）')

doc.add_paragraph()

# 4.2
add_rich_para([('A股投研数据库 —— 智能投研Agent与知识库系统', True)], space_after=0)
add_para('核心贡献：基于Hermes Agent构建A股投研数据自动化Pipeline', size=9, space_after=4)
add_bullet('项目背景：为个人投资者提供A股财务数据的自动化分析工具，支持自然语言查询与知识库检索。')
add_rich_bullet([('系统架构：', True)])
add_bullet('数据抓取层：基于Hermes Agent框架，编写Skill调用财务数据API，自动抓取上市公司财报、估值、行业对比等数据；', level=1)
add_bullet('分析整理层：开发数据处理Skill，对原始财务数据进行清洗、结构化、指标计算，生成结构化的LLM-Wiki知识条目；', level=1)
add_bullet('知识发布层：利用Obsidian Digital Garden将LLM-Wiki发布为静态网站，实现知识的可视化浏览与检索；', level=1)
add_bullet('用户交互层：通过微信公众号与Hermes Agent进行交互，发送自然语言查询，Agent调用对应Skill返回分析结果。', level=1)
add_rich_bullet([('技术亮点：', True), ('Agent Skill机制封装数据源；LLM-Wiki保证知识条目结构化与可追溯性；微信交互实现"对话即查询"。', False)])
add_bullet('技术栈：Python，Hermes Agent框架，FastAPI，Obsidian + Digital Garden，微信公众平台API')

doc.add_paragraph()

# 4.3
add_rich_para([('本地双卡大模型推理环境（基础设施）', True)], space_after=0)
add_para('核心贡献：独立完成本地双GPU推理系统的搭建与模型部署', size=9, space_after=4)
add_bullet('配置双GPU（NVIDIA）推理系统，完成驱动/CUDA环境/PyTorch部署，跑通LLaMA/Mistral/Qwen等主流模型的推理与微调流程。')
add_bullet('深入理解PagedAttention（按需分配KV Cache，与TDVF的Lazy Accept机制一脉相承）、Continuous Batching（动态合并请求，与vTPM的并发调度逻辑相通）等推理加速原理，正在系统学习CUDA Kernel编写与sglang源码架构。')
add_bullet('技术栈：Python，PyTorch，CUDA（基础），Docker，vLLM')

# ============================================================
# AI辅助工程方法论
# ============================================================
add_heading_styled('AI辅助工程方法论（Vibe-Coding实践）', level=2)
add_bullet('产品定义先行：在编写任何代码之前，首先完成产品的功能定义与架构设计，明确系统边界、核心模块划分及关键数据流。')
add_bullet('任务原子化拆解：将复杂项目分解为最小可交付功能集（Atomic Feature Set），与AI进行多轮深度讨论，确保AI明确无误地理解每个任务的目标、约束条件、边界限制及异常处理策略。')
add_bullet('Design Doc驱动开发：要求AI在实现每个最小功能集前，先输出设计文档（Design Doc），经Review通过后方可进入编码阶段。')
add_bullet('代码可理解性审查：AI生成的代码在提交前，必须经过逐行理解与审查，杜绝"黑盒代码"进入代码库。')
add_bullet('核心认知：在AI时代，工程师的核心价值从"编写代码"转向"定义问题、架构决策、质量把关"。')

# ============================================================
# 专业技能
# ============================================================
add_heading_styled('专业技能', level=2)

skills_data = [
    ('系统性能优化方法论', '20年+C/C++经验，精通Lazy/按需加载、Pre-allocation/池化、多核并行化等优化范式；熟练使用Perf/SystemTap/GDB进行微架构级性能分析'),
    ('高并发架构设计', '精通Context隔离、多租户路由、异步IO+事件驱动等架构模式；有将单实例串行服务改造为高并发多租户服务的完整实战经验'),
    ('Rust系统编程', '具备Rust实战经验，理解所有权/生命周期/零成本抽象，熟练使用Tokio异步运行时构建高性能服务'),
    ('AI工程化', '精通Python，具备PyTorch模型加载/推理/后训练实操经验；理解Transformer架构及KV Cache管理机制；可快速上手sglang/vLLM等推理框架'),
    ('硬件底层', '精通X86内存管理、TLB/Cache优化、NUMA亲和性、PCIe/DMA性能调优'),
    ('协议与标准', '深入理解TPM 2.0、SPDM、TDISP等协议，具备极强的英文Spec阅读与标准落地能力'),
]

table = doc.add_table(rows=len(skills_data)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_font(table.rows[0].cells[0], '类别', bold=True)
set_cell_font(table.rows[0].cells[1], '具体能力', bold=True)
for i, (cat, detail) in enumerate(skills_data):
    set_cell_font(table.rows[i+1].cells[0], cat, bold=True)
    set_cell_font(table.rows[i+1].cells[1], detail)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(12.5)

doc.add_paragraph()

# ============================================================
# 语言能力
# ============================================================
add_heading_styled('语言能力', level=2)
add_bullet('中文：母语')
add_bullet('英文：近20年外企工作环境，可作为工作语言，具备流畅的技术文档撰写与Spec阅读能力')

# ============================================================
# Footer
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2026年7月更新')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\work\temp\徐敏_简历.docx'
doc.save(output_path)
print(f'简历已生成: {output_path}')
